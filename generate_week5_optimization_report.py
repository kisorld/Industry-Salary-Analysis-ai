from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reports") / "5.优化报告-行业薪酬洞察AI智能体.docx"


def set_font(run, font_name: str = "Microsoft YaHei", size: float | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def cell_text(cell, text: str, bold: bool = False, color: str = "000000", size: float = 8.8) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_font(run, size=size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def para(doc: Document, text: str, size: float = 10.5, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size=size)
    run.bold = bold


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=10.5)


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, "Consolas", 9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell_width(cell, widths[idx])
        cell_text(cell, header, bold=True, color="FFFFFF", size=9.2)
        shade(cell, "5B9BD5")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell_width(cells[idx], widths[idx])
            cell_text(cells[idx], value, size=8.3 if len(headers) >= 4 else 8.8)
    doc.add_paragraph()


def build() -> Path:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Title", 18, "17324D"),
        ("Heading 1", 15, "2E74B5"),
        ("Heading 2", 12, "1F4D78"),
        ("Heading 3", 11, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《软件开发综合实训》优化报告")
    set_font(run, size=18)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 50, 77)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("（工业软件 + AI赋能课题）")
    set_font(run, size=11)

    for line in [
        "项目名称：行业薪酬洞察 AI 智能体",
        "所属赛道：通用场景赛道",
        "提交阶段：第5周 - BadCase分析与迭代优化",
        "提交日期：2026年6月13日",
    ]:
        para(doc, line)

    heading(doc, "一、Top 3 错误类型与根因分析 (Root Cause Analysis)", 1)
    para(doc, "基于第4周系统测试报告和本周实际联调过程中的Bad Case分析，我们将系统当前最影响可用性与答辩演示稳定性的缺陷归纳为以下三类，并进行了根因剖析：")
    add_table(
        doc,
        ["错误类型", "出现频率", "典型场景", "深度根因分析"],
        [
            [
                "1. 平台采集不稳定",
                "高频（约占Bad Case的35%）",
                "用户选择BOSS直聘或智联招聘，输入岗位关键词和城市后，系统只返回no_response，或只能采集到少量岗位。",
                "根因：平台页面依赖动态接口、登录状态、地区上下文和访问频率。旧版本只监听固定接口（BOSS: joblist；智联: search/positions），当页面未触发接口、进入登录页或验证码页时，系统只能返回空结果，缺少逐页诊断和登录状态识别。",
            ],
            [
                "2. URL解析与网页阅读泛化不足",
                "中高频（约占Bad Case的25%）",
                "用户输入通用入口URL、动态列表页或无薪资详情页时，系统可能返回空结果，或把页面中类似薪资的数字误识别为薪资。",
                "根因：URL采集模块同时承担网页读取、岗位识别和薪资抽取任务。旧规则只要匹配到薪资格式就可能生成岗位记录，缺少“薪资必须与岗位上下文绑定”的校验；AI阅读结果也缺少事实校验，存在模型返回页面中没有的薪资字段的风险。",
            ],
            [
                "3. AI智能体自主决策不足",
                "中频（约占Bad Case的20%）",
                "平台采集失败或有效薪资样本不足时，智能体仍按固定流程继续分析，无法明确判断是否应补采、回退CSV或降低结论强度。",
                "根因：原Agent更接近线性流程编排器，执行顺序为采集→清洗→统计→AI总结→报告。工具返回异常后缺少Observe与Reflect步骤，也没有基于样本量、无薪资样本和平台无响应页进行数据质量判断。",
            ],
        ],
        [1450, 1500, 3100, 4700],
    )

    heading(doc, "二、迭代优化方案与改动细节", 1)
    para(doc, "针对上述根因，我们在本周实施了三轮迭代优化，具体改动如下：")

    heading(doc, "1. 针对“平台采集不稳定”：引入登录检测、逐页诊断与最近任务输出管理", 2)
    para(doc, "改动方案：", bold=True)
    bullet(doc, "平台采集诊断增强：新增 parse_listener_logs()，对BOSS/智联监听日志进行结构化解析，输出 attempted_pages、successful_pages、no_response_pages、empty_pages、page_details、diagnosis 和 advice。")
    bullet(doc, "登录检测机制：在DrissionPage打开平台页面后，先检测当前页面是否进入登录、验证码、扫码、安全验证等状态；若检测到未登录，则返回 login_required。")
    bullet(doc, "自动打开登录页：当 AUTO_OPEN_LOGIN_PAGE=true 时，系统自动打开对应平台登录页；用户登录完成后再次点击采集即可复用 browser_profile 中的登录状态。")
    bullet(doc, "后台采集策略：新增 PLATFORM_SCRAPE_HEADLESS=true 配置，用户已登录时后台抓取和分析，不主动打开浏览器窗口。")
    bullet(doc, "输出管理：新增 outputs/task_runs 最近任务目录，每次智能体或平台分析任务归档CSV、日志、分析JSON和Word报告，并自动只保留最近10次任务。")
    code_block(
        doc,
        "关键文件：backend/app/platform_scrapers.py\n"
        "新增函数：parse_listener_logs(), detect_login_required(), open_platform_login_page()\n"
        "新增配置：PLATFORM_SCRAPE_HEADLESS, AUTO_OPEN_LOGIN_PAGE, BROWSER_USER_DATA_DIR",
    )

    heading(doc, "2. 针对“URL解析与网页阅读泛化不足”：实施薪资绑定分级与AI结果校验", 2)
    para(doc, "改动方案：", bold=True)
    bullet(doc, "薪资绑定分级：新增 salary_binding_level()，将薪资识别结果分为 strong、weak、none 三档。")
    bullet(doc, "强绑定进入统计：岗位标题、公司、城市、学历、经验等上下文与薪资同时出现时，薪资进入 pandas 统计。")
    bullet(doc, "弱绑定进入待确认：存在岗位相关词和疑似薪资，但上下文不完整时，写入 partial_rows，并保存 salary_text_candidate 与 salary_confidence=weak，不参与均值、中位数和分位值计算。")
    bullet(doc, "无绑定不采纳：只有裸薪资数字、公司规模、日期、页脚数字等无岗位上下文的内容不进入统计，降低虚假薪资污染。")
    bullet(doc, "AI阅读后处理：新增 validate_ai_reader_items()，模型返回的薪资必须能在网页可见文本中校验，否则清空 salary_text 并添加 quality_warning。")
    code_block(
        doc,
        "示例：\n"
        "算法工程师 腾讯 深圳 本科 20-35K -> strong，进入统计\n"
        "岗位职责 算法策略 20-35K -> weak，进入待确认样本\n"
        "20-35K -> none，不采纳",
    )

    heading(doc, "3. 针对“AI智能体自主决策不足”：构建反思与补救决策层", 2)
    para(doc, "改动方案：", bold=True)
    bullet(doc, "新增 build_agent_reflection()：根据有效薪资样本数、partial_rows数量、平台无响应页数、失败平台等信息生成反思结果。")
    bullet(doc, "新增执行步骤“反思与补救决策”：在数据汇总后、统计分析前判断当前数据是否足以生成强薪资结论。")
    bullet(doc, "样本不足时降低结论强度：当只有无薪资岗位或有效薪资样本较少时，系统保留岗位信息，但提示不生成强薪资结论。")
    bullet(doc, "采集失败时给出补救路径：提示用户补充关键词、减少页数、切换平台、完成登录或使用CSV/Excel导入。")
    bullet(doc, "分析结果中写入 agent_reflection，报告和前端都可以读取该反思结果。")
    code_block(
        doc,
        "关键文件：backend/app/agent.py\n"
        "新增函数：build_agent_reflection()\n"
        "新增返回字段：reflection、analysis.agent_reflection、task_archive",
    )

    heading(doc, "三、优化前后效果对比与量化分析", 1)
    para(doc, "我们在第4周暴露的典型Bad Case和本周联调场景上重新运行优化后的系统，结果对比如下：")
    add_table(
        doc,
        ["评测指标", "优化前 / Week4", "优化后 / Week5", "变化趋势", "关键归因"],
        [
            ["平台失败可解释性", "仅显示no_response或空结果", "返回逐页诊断、登录状态、失败建议", "显著提升", "parse_listener_logs + login_required诊断"],
            ["未登录场景处理", "后台采集失败，用户不知道原因", "检测到未登录后自动打开登录页", "新增能力", "detect_login_required + open_platform_login_page"],
            ["已登录后台采集", "调试时可能频繁打开浏览器", "已登录时后台抓取和分析", "体验提升", "PLATFORM_SCRAPE_HEADLESS配置"],
            ["虚假薪资误入统计", "可能把无关数字当薪资", "强/弱/无绑定分级，弱绑定不入统计", "风险下降", "salary_binding_level三档判断"],
            ["AI网页阅读可信度", "模型输出缺少事实校验", "薪资必须能在可见文本中校验", "风险下降", "validate_ai_reader_items后处理"],
            ["智能体失败处理", "线性流程，失败后只给最终状态", "新增反思与补救决策步骤", "能力增强", "build_agent_reflection"],
            ["输出文件管理", "旧文件可能长期堆积或被覆盖", "只归档最近10次任务输出", "更规范", "outputs/task_runs + prune_task_runs"],
        ],
        [1900, 2200, 2400, 1300, 2700],
    )

    heading(doc, "典型Bad Case修复示例：", 2)
    para(doc, "Case 1：智联招聘未登录导致采集失败。", bold=True)
    para(doc, "优化前输出：系统返回空数据或 no_response，用户无法判断是关键词错误、平台无结果还是未登录。")
    para(doc, "优化后行为：")
    code_block(
        doc,
        "Observation: diagnostics.login_required = true\n"
        "Action: 自动打开 https://passport.zhaopin.com/login\n"
        "Final: 提示用户完成登录后重新采集；登录状态保存到 backend/.browser_profile。",
    )
    para(doc, "Case 2：页面中出现类似薪资的数字但没有岗位上下文。", bold=True)
    para(doc, "优化前输出：可能将页面中的数字误识别为 salary_text，进入统计。")
    para(doc, "优化后行为：")
    code_block(
        doc,
        "salary_binding_level('20-35K') -> none，不采纳\n"
        "salary_binding_level('岗位职责 算法策略 20-35K') -> weak，进入待确认样本\n"
        "salary_binding_level('算法工程师 腾讯 深圳 本科 20-35K') -> strong，进入统计",
    )
    para(doc, "Case 3：平台采集样本不足但仍生成薪资结论。", bold=True)
    para(doc, "优化前输出：流程继续执行，AI可能给出强结论。")
    para(doc, "优化后行为：")
    code_block(
        doc,
        "Reflect: 有效薪资样本偏少或仅有无薪资岗位。\n"
        "Decision: 保留岗位核心信息，但降低薪资结论强度，并建议补充CSV/Excel或重新采集。",
    )

    heading(doc, "四、遗留问题与后续计划", 1)
    para(doc, "尽管本次迭代解决了三大核心痛点，但仍有长尾问题待解决：")
    bullet(doc, "平台公开页面仍可能受登录、验证码、游客态、地区和平台风控影响；系统只做合规公开采集与失败诊断，不绕过访问控制。")
    bullet(doc, "薪资绑定分级采用规则判断，特殊岗位名称或复杂页面结构仍可能出现弱绑定，需要后续增加更多岗位词典或平台适配器。")
    bullet(doc, "当前任务归档已限制最近10次，但前端尚未提供历史任务管理页面，只能通过 /api/task-runs 查看。")
    bullet(doc, "长任务仍是同步请求，20页采集和AI分析可能超时，后续建议引入异步任务队列和进度轮询。")
    para(doc, "第6周计划：", bold=True)
    bullet(doc, "将平台采集改造成异步任务，前端展示实时进度、当前页数、成功页数和失败原因。")
    bullet(doc, "增加历史任务页面，展示最近10次任务的CSV、分析JSON和Word报告。")
    bullet(doc, "继续扩展岗位词典和平台字段适配器，提升弱绑定样本的自动确认能力。")
    bullet(doc, "完善报告生成模块，将数据质量、采集诊断和智能体反思结果自动写入薪酬分析报告。")

    para(doc, "日期：2026年6月13日")
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
