from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def insight_lines(analysis: dict[str, object], ai_text: str = "") -> list[str]:
    overall = analysis.get("overall", {})
    lines = [
        f"本次数据共 {analysis.get('total', 0)} 条，有效薪资样本 {analysis.get('valid', 0)} 条，解析成功率 {analysis.get('parse_accuracy', 0)}%。",
        f"整体月薪中位数为 {overall.get('中位数K月', 0)}K，P75 为 {overall.get('P75', 0)}K，P90 为 {overall.get('P90', 0)}K。",
    ]
    by_category = analysis.get("by_category") or []
    by_city = analysis.get("by_city") or []
    if by_category:
        top = max(by_category, key=lambda row: row.get("中位数", 0))
        lines.append(f"岗位类别中，{top['维度']} 的中位薪资最高，为 {top['中位数']}K/月。")
    if by_city:
        top = max(by_city, key=lambda row: row.get("中位数", 0))
        lines.append(f"城市维度中，{top['维度']} 的中位薪资最高，为 {top['中位数']}K/月。")
    if ai_text:
        lines.append(ai_text)
    return lines


def add_table(doc: Document, rows: list[dict[str, object]], title: str, max_rows: int = 12) -> None:
    doc.add_heading(title, level=2)
    if not rows:
        doc.add_paragraph("暂无数据。")
        return
    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = str(header)
    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            cells[idx].text = str(row.get(header, ""))


def build_report(analysis: dict[str, object], path: Path, ai_text: str = "", title: str = "行业薪酬洞察报告") -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2"]:
        styles[style_name].font.name = "Microsoft YaHei"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Microsoft YaHei"

    doc.add_heading("一、结论摘要", level=1)
    for line in insight_lines(analysis, ai_text):
        doc.add_paragraph(line)

    doc.add_heading("二、样本与方法", level=1)
    doc.add_paragraph("系统通过 CSV/Excel 导入或公开页面采集 PoC 获取岗位数据，使用 pandas 批量清洗、规则化薪资解析和分位值统计。AI 只负责基于统计结果生成解释，不直接编造数值。")

    doc.add_heading("三、统计结果", level=1)
    add_table(doc, analysis.get("by_category", []), "按岗位类别统计")
    add_table(doc, analysis.get("by_city", []), "按城市统计")
    add_table(doc, analysis.get("by_experience", []), "按经验统计")
    add_table(doc, analysis.get("by_education", []), "按学历统计")

    doc.add_heading("四、风险提示", level=1)
    doc.add_paragraph("样本来自演示数据或公开页面 PoC，若样本量较小，不宜作为企业薪酬决策的唯一依据。真实平台采集需遵守平台规则，不绕过登录、验证码、付费墙或访问限制。")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path

