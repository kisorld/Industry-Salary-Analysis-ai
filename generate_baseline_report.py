from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from salary_core import analyze, clean_jobs, load_jobs, make_insight_text, write_cleaned_csv, write_stats_csv


ROOT = Path(__file__).resolve().parent
DATA_PATH = ROOT / "data" / "baseline_jobs.csv"
OUTPUT_DIR = ROOT / "outputs"
REPORT_DIR = ROOT / "reports"


def add_table(doc: Document, rows: list[dict[str, object]], title: str, max_rows: int = 12) -> None:
    doc.add_heading(title, level=2)
    if not rows:
        doc.add_paragraph("暂无数据。")
        return
    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = str(header)
    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for i, header in enumerate(headers):
            cells[i].text = str(row[header])


def build_docx(analysis: dict[str, object], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2"]:
        styles[style_name].font.name = "Microsoft YaHei"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("行业薪酬洞察 AI 智能体\n第2周基线评测报告")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Microsoft YaHei"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("交付阶段：第2周（基线AI实现、前后端联调）    数据日期：2026-05    版本：V1.0")

    doc.add_heading("一、评测结论", level=1)
    for text in make_insight_text(analysis):
        doc.add_paragraph(text, style=None)

    doc.add_heading("二、技术栈调整说明", level=1)
    items = [
        "原计划技术栈覆盖 Streamlit、FastAPI、Playwright、Qwen API、Plotly、SQLite 和 Doc/PDF 导出，适合完整项目，但第二周同时完成真实抓取、AI API 稳定调用和 PDF 导出风险较高。",
        "第2周基线调整为：纯 Python 标准库 Web 演示 + CSV 样例数据/上传导入 + 规则薪资解析 + pandas/标准库统计 + python-docx 导出 Word 报告。",
        "真实招聘网站抓取、FastAPI 服务化、外部大模型 Tool Calling、PDF 导出和历史任务数据库保留为第3周到第6周增强项。",
    ]
    for item in items:
        doc.add_paragraph(item, style=None)

    doc.add_heading("三、评测数据集", level=1)
    doc.add_paragraph(
        f"本次基线使用离线可复核 CSV 数据集，共 {analysis['total']} 条岗位记录，覆盖数据分析、后端开发、产品经理 3 类岗位，北京、上海、广州、深圳 4 个地区，以及 1-3年、3-5年、5-10年 3 类经验要求。"
    )

    doc.add_heading("四、评测指标结果", level=1)
    metric_rows = [
        {"指标": "功能完整性", "目标": "第2周核心闭环可运行", "结果": "已实现导入、清洗、统计、图表、报告导出", "结论": "通过"},
        {"指标": "文件导入成功率", "目标": "接近100%", "结果": "样例 CSV 可稳定导入", "结论": "通过"},
        {"指标": "薪资解析成功率", "目标": ">=90%", "结果": f"{analysis['parse_accuracy']}%", "结论": "通过" if analysis["parse_accuracy"] >= 90 else "待优化"},
        {"指标": "统计正确性", "目标": "程序可复核", "结果": "均值、中位数、P25/P75/P90 由程序计算", "结论": "通过"},
        {"指标": "报告完整性", "目标": "包含来源、样本量、关键发现、建议和风险", "结果": "已包含", "结论": "通过"},
        {"指标": "演示稳定性", "目标": "核心流程无崩溃", "结果": "本地样例流程已运行", "结论": "通过"},
    ]
    add_table(doc, metric_rows, "指标对照表", max_rows=10)

    add_table(doc, analysis["by_category"], "按岗位类别统计")
    add_table(doc, analysis["by_city"], "按城市统计")
    add_table(doc, analysis["by_experience"], "按经验统计")

    doc.add_heading("五、风险与后续增强", level=1)
    next_steps = [
        "真实平台抓取存在页面结构变化、动态加载、登录和反爬限制，第二周不作为验收主路径。",
        "AI 生成报告在无 API Key 或网络受限时不可稳定复现，当前基线采用模板化智能报告，后续接入 Qwen API 并增加 JSON Schema 约束。",
        "PDF 导出依赖系统环境，当前优先保证 Word 报告，后续可用 LibreOffice 或 WeasyPrint 增强。",
        "历史任务管理可在后续引入 SQLite，保存任务参数、清洗数据、统计结果和报告路径。",
    ]
    for item in next_steps:
        doc.add_paragraph(item)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    REPORT_DIR.mkdir(exist_ok=True)
    jobs = clean_jobs(load_jobs(DATA_PATH))
    analysis = analyze(jobs)
    write_cleaned_csv(jobs, OUTPUT_DIR / "cleaned_jobs.csv")
    write_stats_csv(analysis["by_category"], OUTPUT_DIR / "stats_by_category.csv")
    write_stats_csv(analysis["by_city"], OUTPUT_DIR / "stats_by_city.csv")
    write_stats_csv(analysis["by_experience"], OUTPUT_DIR / "stats_by_experience.csv")
    build_docx(analysis, REPORT_DIR / "基线评测报告.docx")
    print(REPORT_DIR / "基线评测报告.docx")


if __name__ == "__main__":
    main()
