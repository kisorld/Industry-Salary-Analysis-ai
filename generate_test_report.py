from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reports") / "6.测试报告.docx"


def set_east_asia(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text: object, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_east_asia(run)
    run.font.size = Pt(9.2)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell_text(cell, header, True)
        shade(cell, "F2F4F7")
        cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell_text(cells[idx], value)
            cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_east_asia(run)


def para(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_east_asia(run)
    run.font.size = Pt(10.5)


def bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_east_asia(run)
    run.font.size = Pt(10.5)


def build() -> Path:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Title", 18, "17324D"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《软件开发综合实训》测试报告")
    set_east_asia(run)
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 50, 77)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("（工业软件 + AI赋能课题）")
    set_east_asia(run)
    run.font.size = Pt(11)

    for line in [
        "项目名称：行业薪酬洞察 AI 智能体",
        "所属赛道：通用场景赛道",
        "提交阶段：第4周 - 系统测试与评测",
        "提交日期：2026年6月7日",
    ]:
        para(doc, line)

    heading(doc, "一、测试概述", 1)
    heading(doc, "测试目标", 2)
    for item in [
        "验证系统是否能够完成“公开招聘信息采集 → CSV/Excel导入 → 字段标准化 → 薪资统计 → AI分析 → Word报告生成”的端到端链路。",
        "重点测试 BOSS直聘、智联招聘两个平台脚本式接口监听采集能力，包括公开搜索URL生成、前20页监听爬取、CSV输出和失败日志记录。",
        "验证 Qwen AI核心、RAG知识库、pandas统计、Vue3可视化和一键启动功能在集成后的稳定性。",
        "评估系统在动态页面、无薪资岗位、空URL、平台无响应、模型超时等边界情况下的鲁棒性。",
    ]:
        bullet(doc, item)

    heading(doc, "测试环境", 2)
    add_table(
        doc,
        ["类别", "环境/工具", "说明"],
        [
            ["操作系统", "Windows 10/11", "本地开发与演示环境"],
            ["前端", "Vue3 + Vite + axios + ECharts", "负责任务输入、结果展示、图表可视化和报告下载"],
            ["后端", "Python FastAPI + uvicorn", "负责接口编排、数据处理、爬虫、RAG和报告生成"],
            ["爬虫工具", "DrissionPage、Playwright、Scrapling", "DrissionPage作为BOSS/智联脚本式监听核心，Playwright/Scrapling作为兜底"],
            ["AI模型", "Qwen OpenAI-compatible API", "用于薪酬建议、URL阅读结果结构化和报告摘要"],
            ["数据分析", "pandas、numpy", "薪资解析、分位值统计、维度聚合和CSV输出"],
            ["RAG", "ChromaDB + sentence-transformers", "用于薪资口径、合规说明和报告建议增强"],
        ],
        [1.3, 2.0, 3.2],
    )

    heading(doc, "二、系统性评测集构建", 1)
    para(doc, "本阶段测试集围绕系统真实业务链路构建，覆盖平台爬取、导入数据、薪资解析、AI报告和异常场景。由于招聘平台公开页面存在游客可见范围、登录、验证码和风控差异，测试重点放在公开可见数据的稳定处理和失败可解释性上。")
    add_table(
        doc,
        ["数据集类别", "样本数量/规模", "构建方式", "核心目的"],
        [
            ["基础样例数据集", "180条", "系统内置 baseline_jobs.csv", "验证CSV/Excel导入、字段标准化和薪资统计能力"],
            ["平台爬取测试集", "BOSS/智联各最多20页", "用户输入岗位关键词、城市，系统合成公开URL并监听接口", "验证平台专项爬虫和CSV输出能力"],
            ["公开URL测试集", "10+ URL", "腾讯招聘详情页、BOSS列表页、智联列表页、普通静态页面", "验证URL解析、动态读取和AI阅读补充能力"],
            ["薪资格式测试集", "多类型格式", "15-20K、20-30K·14薪、200-300元/天、30-50万年薪、面议", "验证薪资解析与异常标记"],
            ["边界/鲁棒集", "20+ 场景", "空URL、无薪资页面、接口无响应、模型超时、脚本壳文本、平台限制", "验证系统不崩溃并返回可解释日志"],
        ],
        [1.35, 1.25, 2.35, 2.0],
    )

    heading(doc, "三、测试结果记录与分析", 1)
    heading(doc, "1. 功能与性能测试汇总", 2)
    add_table(
        doc,
        ["评测维度", "核心指标", "测试结果", "达标情况", "分析"],
        [
            ["数据获取", "CSV/Excel导入", "可正常导入样例和上传文件", "达标", "字段映射与缺失列补全逻辑稳定"],
            ["平台采集", "BOSS接口监听", "支持监听 joblist 并抽取岗位字段", "基本达标", "受平台公开访问状态影响，未监听到接口时记录 no_response"],
            ["平台采集", "智联接口监听", "支持监听 search/positions 并抽取岗位字段", "基本达标", "已修复 listen.wait 返回 False 导致后端崩溃的问题"],
            ["URL生成", "按平台/城市/岗位生成公开URL", "可展示生成URL", "达标", "前端新增“已生成公开采集URL”区域"],
            ["结构化处理", "字段标准化", "岗位、公司、城市、薪资、经验、学历、来源可统一字段", "达标", "平台字段已映射到 job_title、salary_text 等统一字段"],
            ["薪资分析", "均值/中位数/P25/P75/P90", "可计算并图表展示", "达标", "仅有效薪资进入统计，无薪资样本进入 partial_rows"],
            ["AI能力", "Qwen报告生成", "可生成AI建议，失败时local_fallback", "达标", "RAG上下文注入增强数据口径说明"],
            ["可视化", "薪资和需求趋势图", "支持岗位、城市、需求数量变化图", "达标", "ECharts图表可随分析结果更新"],
            ["报告导出", "Word报告", "可生成并下载报告", "达标", "输出 agent_salary_insight_report.docx 或 platform_salary_report.docx"],
            ["稳定性", "异常不崩溃", "核心异常已加保护", "基本达标", "平台接口无响应仍需继续优化提示和重试策略"],
        ],
        [1.0, 1.4, 1.7, 0.8, 2.2],
    )

    heading(doc, "2. 典型Bad Case深度分析", 2)
    heading(doc, "Bad Case 1：智联监听接口未返回响应导致后端500", 3)
    para(doc, "输入：用户选择智联招聘，填写岗位关键词后点击“脚本爬取并分析”。")
    para(doc, "系统行为：DrissionPage 的 listen.wait() 在未监听到 search/positions 接口时返回 False，旧代码继续读取 resp.response.body，导致 AttributeError: bool object has no attribute response。")
    para(doc, "根因分析：接口监听逻辑假设每次都能获取响应对象，没有处理超时、无响应或页面未触发接口的情况。")
    para(doc, "改进结果：新增 response_body_or_none()，当监听结果为空或为 False 时记录 zhaopin_listener_page:N:no_response，不再导致 FastAPI 崩溃。")

    heading(doc, "Bad Case 2：BOSS通用空URL无法分析出岗位", 3)
    para(doc, "输入：https://www.zhipin.com/web/geek/jobs。")
    para(doc, "系统行为：页面没有 query、city、position 等明确条件，可能只显示默认入口、登录提示、城市选择或前端壳，岗位卡片数量为0或极少。")
    para(doc, "根因分析：该URL不是稳定的数据入口，平台依赖用户状态、城市上下文和前端异步接口。")
    para(doc, "改进方向：前端要求用户输入岗位关键词和城市，系统合成带 query/city 参数的公开URL，并展示生成结果供用户确认。")

    heading(doc, "Bad Case 3：复杂列表页只返回少量岗位", 3)
    para(doc, "输入：BOSS或智联部分列表页。")
    para(doc, "系统行为：即使滚动页面，也可能只返回少量岗位或脚本壳文本。")
    para(doc, "根因分析：平台可能对游客态限制展示数量，或使用虚拟列表、登录校验、验证码、风控策略。系统不能也不应绕过这些访问控制。")
    para(doc, "改进方向：当前采用DrissionPage接口监听作为优先方案，并保留DOM采集、Playwright和Scrapling兜底；后续可增加更多公开数据源或人工CSV导入。")

    heading(doc, "Bad Case 4：大模型抽取阶段超时", 3)
    para(doc, "输入：包含大量导航、城市筛选、脚本文本的网页内容。")
    para(doc, "系统行为：后端已读取文本，但Qwen API调用出现 read operation timed out。")
    para(doc, "根因分析：传入模型的无关文本过多，模型处理时间超过前端或后端超时限制。")
    para(doc, "改进结果：系统增加岗位卡片优先输入、focus_job_text正文聚焦、120秒模型超时和脚本壳过滤。")

    heading(doc, "四、测试结论与系统状态", 1)
    heading(doc, "整体评价", 2)
    para(doc, "行业薪酬洞察AI智能体已完成从数据获取、结构化处理、薪资统计、AI分析、可视化到Word报告生成的主链路。系统具备按平台、地区、岗位关键词生成公开URL并进行脚本式爬取的能力，也支持CSV/Excel导入作为低风险数据来源。")
    para(doc, "相比基线版本，当前系统已从单一薪资分析工具升级为具备工具调用、平台采集、RAG增强和报告生成能力的任务编排型AI智能体。")

    heading(doc, "遗留问题与风险", 2)
    for item in [
        "BOSS直聘、智联招聘等平台的公开页面可见数据受登录、验证码、地区、游客态和风控影响，不能保证任意条件均可采集到20页数据。",
        "系统不会绕过登录、验证码、付费墙或平台风控，因此受限页面需要通过CSV/Excel导入或合规数据源补充。",
        "AI生成的建议依赖有效薪资样本量，样本过少时只能作为演示和辅助参考，不能作为企业薪酬决策唯一依据。",
        "前端任务执行为同步请求，20页爬取和AI分析可能耗时较长，后续建议改为异步任务队列和进度轮询。",
    ]:
        bullet(doc, item)

    heading(doc, "交付建议", 2)
    for item in [
        "当前系统已具备Demo演示条件，建议演示时优先使用带岗位关键词和城市参数的公开URL，避免空入口URL。",
        "答辩时重点展示“生成公开采集URL → 脚本监听爬取 → CSV输出 → pandas统计 → AI报告”的完整链路。",
        "下一阶段建议增加异步任务队列、爬取进度条、采集失败诊断面板和更多公开数据源适配。",
        "建议继续强化CSV/Excel导入能力，使系统在平台限制场景下仍能稳定完成薪资分析目标。",
    ]:
        bullet(doc, item)

    para(doc, "日期：2026年6月7日")
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
