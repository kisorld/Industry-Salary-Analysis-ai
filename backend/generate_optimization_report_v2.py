from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "4.优化对比报告-采集RAG增强版.docx"


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1
    for name, size, color in [
        ("Title", 18, "0B2545"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Microsoft YaHei"
    run.font.color.rgb = RGBColor.from_string("0B2545")


def add_meta(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[object]]) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_report() -> Path:
    doc = Document()
    setup_styles(doc)

    add_title(doc, "《软件开发综合实训》优化对比报告")
    add_meta(
        doc,
        [
            "（工业软件 + AI赋能课题）",
            "项目名称：行业薪酬洞察 AI 智能体",
            "所属赛道：通用场景赛道",
            "提交阶段：第4周 - 采集、RAG与系统联调增强",
            "提交日期：2026年5月31日",
        ],
    )

    doc.add_heading("一、优化概览", level=1)
    doc.add_heading("优化目标回顾", level=2)
    doc.add_paragraph(
        "基于第2周《基线评测报告》和第3周《优化对比报告》暴露的问题，本阶段重点优化真实公开页面采集、URL解析、网页数据抽取、提示词约束、RAG检索增强和前后端一键启动能力。优化目标不是绕过招聘平台限制，而是在合法公开页面范围内提升采集稳定性、可诊断性和报告生成可靠性。"
    )
    doc.add_heading("增强方案实施摘要", level=2)
    add_bullets(
        doc,
        [
            "采集 PoC 增强：由单一文本窗口匹配升级为 JSON-LD、article/li 列表块、职位卡片 class/id 与上下文窗口的组合抽取。",
            "URL 解析增强：支持自动补全协议、多 URL 粘贴、逗号/分号/换行分隔、URL合法性校验和失败原因返回。",
            "网页爬取能力增强：新增访问状态、页面大小、文本节点数、候选职位块数、命中岗位数等采集日志。",
            "提示词优化：将模型输出固定为“样本说明、关键发现、企业建议、风险提示、数据口径”，并明确禁止编造统计值和工具参数。",
            "RAG 功能：新增本地薪酬口径知识库，报告生成前自动检索薪资口径、样本量解释、数据合规、岗位归一和报告建议。",
            "一键启动：新增 start_all.bat，自动检查后端虚拟环境和前端依赖，并分别启动 FastAPI 与 Vue3/Vite。",
        ],
    )

    doc.add_heading("二、“基线失败 vs 增强成功”典型案例对比", level=1)
    cases = [
        (
            "案例 1：URL 输入不规范",
            "基线失败：用户输入 example.com/jobs 或一次粘贴多个链接时，后端只接受单个完整 URL，容易报错或无法识别。",
            "增强成功：新增 normalize_url 与 resolve_url_inputs，自动补全 https://，支持换行、逗号、分号分隔的多 URL，并对非法 URL 返回明确错误。",
            "成功分析：URL 解析从“前端必须输入标准格式”升级为“后端负责容错和标准化”，降低演示和真实使用门槛。",
        ),
        (
            "案例 2：网页结构化岗位块识别",
            "基线失败：采集器只按文本节点附近窗口匹配，容易把 3-5年 识别为薪资，或将前一个岗位标题和后一个薪资错配。",
            "增强成功：采集器先识别 JSON-LD、article/li、职位卡片 class/id，再使用上下文窗口兜底；并修正薪资正则，要求薪资范围带 K/千/万/元等单位。",
            "成功分析：结构化块优先策略减少字段错配，薪资正则收紧后避免将经验年限误判为薪资。",
        ),
        (
            "案例 3：AI 报告口径不稳定",
            "基线失败：AI 报告只接收统计摘要，缺少薪资口径、样本量边界和数据合规说明，容易生成泛化建议。",
            "增强成功：新增 RAG 知识库，报告生成前检索“薪资口径、样本量解释、数据合规、岗位归一、报告建议”等片段并注入 Prompt。",
            "成功分析：RAG 让模型回答有明确依据，即使 Qwen API 不可用，本地 fallback 报告也会展示命中的知识库口径。",
        ),
        (
            "案例 4：前后端启动步骤复杂",
            "基线失败：用户需要分别进入 backend 和 frontend，手动启动 uvicorn 与 npm run dev，容易出现目录错误、依赖未安装或端口混乱。",
            "增强成功：根目录新增 start_all.bat，自动检查 backend/.venv 与 frontend/node_modules，并在两个窗口分别启动后端和前端。",
            "成功分析：一键启动提升演示可靠性，尤其适合答辩现场或非开发成员复现实验。",
        ),
    ]
    for title, baseline, enhanced, analysis in cases:
        doc.add_heading(title, level=2)
        doc.add_paragraph("【基线失败案例】")
        doc.add_paragraph(baseline)
        doc.add_paragraph("【增强成功案例】")
        doc.add_paragraph(enhanced)
        doc.add_paragraph(analysis)

    add_table(
        doc,
        "三、增强后性能变化量化分析",
        ["评测指标", "基线/上一版", "增强后", "变化趋势", "原因分析"],
        [
            ["URL解析能力", "仅单 URL，需完整协议", "支持协议补全、多 URL、合法性校验", "显著提升", "新增 url_tools 标准化与批量解析"],
            ["采集字段错配风险", "文本窗口匹配，可能错配", "JSON-LD + article/li + 职位卡片 + 上下文兜底", "降低", "结构化块优先，窗口兜底"],
            ["薪资误识别", "可能把 3-5年 识别为薪资", "薪资范围必须带薪资单位或薪资上下文", "降低", "收紧 SALARY_RE 规则"],
            ["采集可诊断性", "失败时只提示未识别", "返回状态、页面大小、文本节点、候选块、命中数", "提升", "新增采集日志 logs"],
            ["RAG能力", "无", "本地知识库检索并注入 Prompt", "新增", "新增 rag.py 与 docs/rag_knowledge_base.md"],
            ["一键启动", "手动分别启动", "start_all.bat 自动检查并启动", "提升", "降低环境和路径错误"],
        ],
    )

    doc.add_heading("四、新错误的产生与深度分析", level=1)
    sections = [
        (
            "1. 采集召回增加后的弱相关匹配",
            "现象：采集器支持更多结构后，可能从导航、推荐岗位或页面底部相似职位中抽取弱相关岗位。根因是公开页面结构差异较大，单靠 HTML 语义标签无法完全判断主内容区。应对策略是在后续加入域名级解析适配器和主内容区过滤规则。",
        ),
        (
            "2. RAG 上下文过度注入",
            "现象：当检索词较泛时，RAG 可能返回“报告建议”等通用片段，模型回答会偏向模板化。根因是当前 RAG 使用轻量词重叠评分，没有向量语义检索和 rerank。应对策略是增加最小相关性阈值、人工维护知识库标题关键词，并在后续引入向量检索。",
        ),
        (
            "3. Qwen Tool Calling 模型兼容性",
            "现象：如果模型名选择不支持工具调用的模型，系统会回退到本地模板报告。根因是不同 Qwen 模型对 OpenAI-compatible tool calling 的支持程度不同。应对策略是在 README 中要求使用控制台确认支持工具调用的模型，并保留 local_fallback 保障演示。",
        ),
        (
            "4. 一键启动依赖本机环境",
            "现象：如果未安装 Python 3.11、Node.js LTS 或未创建 backend/.venv，一键启动会提示错误而不是自动安装全部环境。根因是系统级软件安装需要用户权限和网络。应对策略是脚本给出明确安装命令，避免静默失败。",
        ),
    ]
    for heading, body in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)

    doc.add_heading("五、下一阶段工作重点", level=1)
    add_bullets(
        doc,
        [
            "引入域名级采集适配器，为常见公开招聘页面分别维护字段抽取规则。",
            "增加采集回归测试集，覆盖 JSON-LD、列表卡片、纯文本页面、无薪资页面和错误 URL。",
            "为 RAG 增加相关性阈值、检索日志和可编辑知识库管理界面。",
            "增加模型输出数值一致性检查，确保报告中的数值全部来自 pandas 统计结果或工具返回。",
            "将采集日志、RAG 命中、Tool Calling 过程在前端分步展示，提高答辩可解释性。",
        ],
    )

    doc.add_heading("六、是否需要新增软件", level=1)
    doc.add_paragraph(
        "本轮优化没有强制引入新的 Python 第三方库，后端继续使用 FastAPI、pandas、python-docx 等现有依赖；前端继续使用 Vue3、Vite、axios、echarts。"
    )
    add_bullets(
        doc,
        [
            "必须：Python 3.11，用于后端虚拟环境。",
            "必须：Node.js LTS，用于 Vue3/Vite 前端。",
            "可选：LibreOffice，用于将 Word 报告渲染为 PDF/PNG 做视觉检查。",
            "可选：Playwright，仅当后续需要采集公开动态渲染页面时安装；当前版本未强制需要。",
        ],
    )
    doc.add_paragraph("日期：2026年5月31日")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_report())
