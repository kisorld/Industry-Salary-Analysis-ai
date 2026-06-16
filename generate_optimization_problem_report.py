from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reports") / "5.优化问题与根因分析报告.docx"


def set_font(run, font_name: str = "Microsoft YaHei", size: float | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def cell_text(cell, text: str, bold: bool = False, color: str = "000000", size: float = 8.8) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_font(run, size=size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph(doc: Document, text: str, size: float = 10.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size=size)


def heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=10.5)


def add_top3_table(doc: Document) -> None:
    headers = ["错误类型", "出现频率", "典型场景", "深度根因分析", "本轮优化措施"]
    rows = [
        [
            "1. 平台采集不稳定",
            "高频（约占 Bad Case 的35%）",
            "用户选择 BOSS直聘或智联招聘，输入岗位关键词和城市后，系统有时只能采集少量岗位，或监听接口返回 no_response。",
            "根因：公开招聘平台页面依赖动态接口、登录状态、地区上下文和访问频率。旧逻辑只记录原始 listener 日志，缺少逐页成功率、无响应页数和可理解失败原因。",
            "已优化：新增 parse_listener_logs()，标准化输出 attempted_pages、successful_pages、no_response_pages、empty_pages、page_details、diagnosis 和 advice；scrape_platform_to_csv() 捕获脚本异常并返回诊断，不再只给出空结果。",
        ],
        [
            "2. URL解析与网页阅读泛化不足",
            "中高频（约占 Bad Case 的25%）",
            "用户输入通用入口URL、动态列表页或无薪资详情页时，系统可能返回空结果、只返回少量岗位，或将页面中的无关数字误识别为薪资。",
            "根因：URL采集模块同时承担网页读取、岗位识别和薪资抽取。旧规则只要识别到薪资格式就可能生成岗位记录，缺少“薪资必须与岗位上下文绑定”的约束；AI阅读也缺少结果校验。",
            "已优化：新增 salary_bound_to_job()，要求薪资必须出现在岗位相关文本窗口中；动态卡片抽取也使用同一校验；AI阅读新增 validate_ai_reader_items()，模型返回的薪资若不在可见文本中出现，会被清空并标记 quality_warning。",
        ],
        [
            "3. AI智能体自主决策不足",
            "中频（约占 Bad Case 的20%）",
            "智能体按固定顺序执行平台采集、URL采集、样例数据、pandas分析和AI总结。采集失败或样本不足时，不能主动说明是否应该补采、回退或降低结论强度。",
            "根因：当前Agent更接近流程编排器，缺少 Plan-Act-Observe-Reflect 闭环。工具返回异常后没有统一反思层，也没有基于样本量、无薪资记录和平台无响应页来判断数据是否足以生成强结论。",
            "已优化：新增 build_agent_reflection()，根据有效薪资样本、partial_rows、平台无响应页和use_sample配置生成反思结论；run_salary_agent()新增“反思与补救决策”步骤，并把 reflection 写入返回结果和 analysis。",
        ],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [1450, 1650, 3000, 3900, 4000]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell_width(cell, widths[idx])
        cell_text(cell, header, bold=True, color="FFFFFF", size=9.3)
        shade(cell, "5B9BD5")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell_width(cells[idx], widths[idx])
            cell_text(cells[idx], value, size=8.2 if idx in {3, 4} else 8.6)


def add_before_after_table(doc: Document) -> None:
    headers = ["优化项", "优化前", "优化后", "涉及文件"]
    rows = [
        [
            "平台采集诊断",
            "只返回原始logs，用户只能看到 no_response 或空rows。",
            "返回逐页诊断、成功页数、无响应页数、诊断原因和操作建议。",
            "backend/app/platform_scrapers.py\nbackend/app/main.py",
        ],
        [
            "URL薪资误识别控制",
            "只要页面出现类似薪资格式的数字，就可能进入候选岗位。",
            "薪资必须与岗位文本绑定；无岗位上下文的薪资数字不进入统计。",
            "backend/app/scraper.py",
        ],
        [
            "AI阅读结果校验",
            "模型返回结果直接展示，缺少事实校验。",
            "模型薪资字段必须能在可见文本中校验，否则清空并给出quality_warning。",
            "backend/app/url_reader.py",
        ],
        [
            "智能体反思决策",
            "固定流程执行，失败后只给最终状态。",
            "新增反思步骤，判断样本不足、无薪资、平台无响应，并给出补救建议。",
            "backend/app/agent.py",
        ],
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [2200, 3600, 4200, 3000]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell_width(cell, widths[idx])
        cell_text(cell, header, bold=True, color="FFFFFF", size=9.3)
        shade(cell, "5B9BD5")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell_width(cells[idx], widths[idx])
            cell_text(cells[idx], value, size=8.6)


def build() -> Path:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Title", 18, "17324D"),
        ("Heading 1", 15, "2E74B5"),
        ("Heading 2", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《软件开发综合实训》优化报告")
    set_font(run, size=18)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 50, 77)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("项目名称：行业薪酬洞察 AI 智能体｜阶段：第5周前三类Bad Case优化")
    set_font(run, size=10.5)

    paragraph(
        doc,
        "本报告基于上一版《优化问题与根因分析报告》中的前三类高优先级问题进行优化复盘。"
        "本轮优化聚焦平台采集诊断、URL解析泛化与AI智能体反思决策，不包含后两类前端流程和数据质量面板重构。",
    )

    heading(doc, "一、Top 3 错误类型、根因与优化措施", 1)
    add_top3_table(doc)

    heading(doc, "二、优化前后效果对比", 1)
    add_before_after_table(doc)

    heading(doc, "三、验证结果", 1)
    for item in [
        "已通过 Python 语法检查：platform_scrapers.py、scraper.py、url_reader.py、agent.py、main.py。",
        "已通过既有薪资解析测试：backend/tests/test_salary_parser.py。",
        "已验证 build_agent_reflection() 在无薪资样本和平台无响应场景下能返回 partial 等级、反思结论和无响应页统计。",
        "由于平台采集依赖真实公开网页和本地浏览器状态，本轮未在报告中承诺任意平台均可稳定采满20页。",
    ]:
        bullet(doc, item)

    heading(doc, "四、遗留问题与后续计划", 1)
    for item in [
        "BOSS直聘、智联招聘公开页面仍可能受登录、验证码、游客态和平台风控影响；系统只做合规公开采集与失败诊断，不绕过访问控制。",
        "后续建议继续优化数据质量面板，展示字段缺失率、平台成功率、无薪资样本和薪资解析失败原因TOP榜。",
        "前端仍建议改造成“输入条件 → 获取数据 → 分析结果 → 生成报告”的阶段式流程，并加入长任务进度提示。",
    ]:
        bullet(doc, item)

    paragraph(doc, "日期：2026年6月13日")
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
