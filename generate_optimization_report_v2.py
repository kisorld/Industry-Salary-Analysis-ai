from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reports") / "5.优化对比报告-复杂URL与AI阅读增强版.docx"


def set_east_asia(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_east_asia(run)
    run.font.size = Pt(9.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, True)
        set_cell_shading(cell, "D9EAF7")
        cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
            cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_east_asia(run)


def add_para(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_east_asia(run)
    run.font.size = Pt(10.5)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_east_asia(run)
    run.font.size = Pt(10.5)


def add_case(doc: Document, title: str, baseline: str, enhanced: str, analysis: str) -> None:
    add_heading(doc, title, 2)
    add_para(doc, "【基线失败案例】")
    add_para(doc, baseline)
    add_para(doc, "【增强成功案例】")
    add_para(doc, enhanced)
    add_para(doc, "成功分析：" + analysis)


def build() -> Path:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Title", 18, "17324D"),
        ("Heading 1", 15, "17324D"),
        ("Heading 2", 12.5, "1F7A8C"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《软件开发综合实训》优化对比报告")
    run.bold = True
    set_east_asia(run)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(23, 50, 77)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("（工业软件 + AI赋能课题）")
    set_east_asia(run)
    run.font.size = Pt(11)

    for line in [
        "项目名称：行业薪酬洞察 AI 智能体",
        "所属赛道：通用场景赛道",
        "提交阶段：第3周/当前阶段 - 复杂URL采集、AI阅读与RAG增强",
        "提交日期：2026年5月31日",
    ]:
        add_para(doc, line)

    add_heading(doc, "一、优化概览", 1)
    add_heading(doc, "优化目标回顾", 2)
    add_para(
        doc,
        "基于第2周《基线评测报告》中暴露的基线能力缺口，以及参考第3周《优化对比报告》的写法，本阶段围绕行业薪酬洞察场景完成了从最小可行系统到增强型演示系统的升级。优化重点不是追求完美覆盖所有招聘平台，而是在公开网页和合法访问边界内，提高系统对复杂URL、动态页面、无薪资页面、岗位列表页和AI结构化抽取的处理能力。",
    )
    add_para(
        doc,
        "基线阶段已经打通 Vue3 前端、Python FastAPI 后端、Qwen API、pandas 批量薪资分析和Word报告生成链路，但在真实公开页面采集、URL解析、模型输出约束、领域知识注入和演示启动稳定性方面仍存在明显短板。本阶段的核心目标是提升采集可解释性、增强大模型阅读网页的能力，并使前后端联调和答辩演示更加稳定。",
    )

    add_heading(doc, "增强方案实施摘要", 2)
    for item in [
        "前后端拆分完善：前端继续使用 Vue3/Vite，后端使用 FastAPI；新增 /api/url/read、/api/url/read/doubao、/api/rag/search、/api/rag/rebuild 等能力接口。",
        "Qwen AI核心增强：保留 Qwen OpenAI-compatible API 调用，支持 Tool Calling 报告生成、RAG上下文注入和本地模板fallback。",
        "复杂URL采集增强：引入 Scrapling DynamicFetcher/StealthyFetcher、Playwright Headless Chromium 和 urllib 兜底，形成多层抓取链路。",
        "AI阅读URL能力：后端先读取动态渲染后的可见文本或结构化详情，再交给 Qwen 抽取岗位核心信息；同时预留豆包URL阅读接口用于对比实验。",
        "腾讯招聘详情页适配：对 careers.tencent.com/jobdesc.html?postId=... 优先提取 postId 并调用公开岗位详情接口，避免被通用列表页滚动逻辑干扰。",
        "岗位列表页增强：增加模拟人类阅读流程，包括正常Chrome UA、中文语言、Referer、等待networkidle、滚动页面、尝试点击展开/查看更多、持续收集岗位卡片。",
        "无薪资页面处理：将结果拆分为 success、partial、failed 三类；有岗位但无薪资时展示 partial_rows，不进入薪资统计，避免误判为采集失败。",
        "RAG增强：支持 ChromaDB + sentence-transformers 向量RAG，默认使用 paraphrase-multilingual-MiniLM-L12-v2，并保留关键词检索fallback。",
        "一键启动：保留 start_all.bat / start_demo.bat，降低答辩现场启动复杂度。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "二、“基线失败 vs 增强成功”典型案例对比", 1)
    add_case(
        doc,
        "案例 1：动态招聘详情页只能返回前端空壳",
        "基线失败：访问腾讯招聘详情页时，静态 urllib 只能拿到约2KB的HTML前端壳，text_nodes 很少，job_like_blocks 为0，系统无法解析岗位名称、地点和职责，最终显示采集失败。",
        "增强成功：系统识别 careers.tencent.com/jobdesc.html?postId=... 后，优先提取 postId 调用腾讯岗位详情接口；若该接口可访问，则得到结构化岗位详情，并交给 Qwen 输出岗位核心信息。",
        "详情页从“通用网页猜测”升级为“站点详情适配 + AI结构化抽取”。对于没有薪资的官网岗位，系统返回 partial 或 salary_text 为空，避免编造薪资。",
    )
    add_case(
        doc,
        "案例 2：页面没有薪资时被误判为采集失败或误造薪资",
        "基线失败：页面没有薪资字段时，系统只按是否识别到 salary_text 判断成功，导致有岗位但无薪资的页面被当作失败；同时宽松薪资正则可能把岗位ID、日期或数字+W误判为薪资。",
        "增强成功：后端将结果拆分为 success、partial、failed。对腾讯详情页只从明确薪资字段读取薪资；如果页面没有公开薪资，salary_text 留空，missing_reason 标注“公开详情页未提供薪资字段”。",
        "系统能区分“采不到岗位”和“岗位存在但无薪资”，并保证薪资统计只使用真实薪资字段，提升评测可信度。",
    )
    add_case(
        doc,
        "案例 3：招聘列表页只能读取少量岗位",
        "基线失败：通用文本窗口只能读取当前首屏或少数DOM节点，BOSS直聘等列表页经常只能返回2条岗位，且Scrapling可能抓到脚本源码而非可见岗位文本。",
        "增强成功：新增脚本壳过滤规则，发现 var staticPath、_PAGE、captcha、\\\\u003C 等脚本特征后自动回退；Playwright端新增滚动到页面底部、持续收集岗位卡片、记录 collect_round 和 job_card_count 的增量采集过程。",
        "列表页从一次性读取升级为“持续滚动 + 增量收集 + 最多80条卡片 + 最多40条交给Qwen抽取”，提高了岗位召回率，并通过日志解释为什么仍可能只返回少数岗位。",
    )
    add_case(
        doc,
        "案例 4：AI阅读URL依赖模型自身联网能力不稳定",
        "基线失败：直接把 URL 发给大模型并要求其动态渲染网页时，如果API通道没有浏览器工具，模型无法真正打开网页，只能根据URL和日志猜测，结果不稳定。",
        "增强成功：改为后端负责动态渲染和读取网页，Qwen负责结构化理解；另新增豆包URL阅读接口，用于验证具备网页阅读能力模型的效果。",
        "职责边界更加清晰：浏览器自动化模拟访问和阅读公开网页，大模型负责理解和抽取，不再依赖Prompt让模型凭空获得浏览器能力。",
    )
    add_case(
        doc,
        "案例 5：领域口径和报告解释不足",
        "基线失败：AI报告只接收 pandas 统计摘要，缺少薪资口径、样本量边界、岗位归一和公开页面合规说明，报告容易泛化。",
        "增强成功：接入本地知识库和向量RAG，报告生成前检索薪资口径、样本量解释、数据合规、岗位归一和报告建议，并注入Prompt。",
        "RAG使报告更贴合本项目的数据口径，即使API不可用，也能通过本地fallback保留基本报告结构和风险提示。",
    )

    add_heading(doc, "三、增强后性能变化量化分析", 1)
    add_para(
        doc,
        "以下指标基于当前系统功能和演示测试情况整理。由于不同招聘网站存在登录、验证码、反爬、地区和游客可见范围差异，公开URL采集指标用于说明优化趋势，而非承诺对所有网站完全覆盖。",
    )
    add_table(
        doc,
        ["评测指标", "基线系统", "增强系统", "变化趋势", "原因分析"],
        [
            ["URL输入兼容性", "单URL、格式要求高", "支持协议补全、多URL分隔、URL校验", "提升", "resolve_url_inputs 与 normalize_url 降低输入错误率"],
            ["动态详情页识别", "静态HTML空壳，无法解析", "腾讯详情页postId接口 + Playwright动态渲染", "显著提升", "对典型详情页建立专用适配器"],
            ["无薪资页面处理", "常被判定失败或误判薪资", "partial状态，岗位展示但不进统计", "显著提升", "success/partial/failed 三状态分流"],
            ["复杂列表页召回", "通常只能读取少量首屏文本", "滚动采集、展开按钮、岗位卡片批量抽取", "提升", "最多收集80条卡片，最多40条交给Qwen"],
            ["AI输出可靠性", "可能泛化或编造字段", "结构化JSON Prompt，禁止编造薪资", "提升", "明确字段缺失时留空并写limits"],
            ["RAG能力", "无领域知识或仅模板说明", "向量RAG + 关键词fallback", "新增", "ChromaDB与本地嵌入模型增强口径说明"],
            ["前端可解释性", "只展示成功/失败", "展示logs、reader_meta、partial_rows、AI阅读结果", "提升", "便于定位是抓取失败、渲染失败还是模型抽取失败"],
            ["响应时间", "较短但能力弱", "动态抓取和AI阅读更慢", "增加", "Playwright/Scrapling和Qwen抽取增加耗时，需要超时控制"],
        ],
        [1.25, 1.35, 1.65, 0.75, 1.5],
    )

    add_heading(doc, "四、新错误的产生与深度分析", 1)
    for title, body in [
        ("1. Scrapling 抓到脚本源码而非可见文本", "现象：在部分BOSS直聘页面中，Scrapling返回大量 var staticPath、_PAGE、captcha、\\\\u003C 等脚本配置，Qwen无法抽取岗位。根因是复杂页面将大量状态数据写入脚本，通用 ::text 提取会误把脚本当正文。应对策略：新增脚本壳过滤和HTML去script/style清洗，疑似脚本时回退到Playwright。"),
        ("2. 列表页可见岗位数量仍受平台限制", "现象：即使滚动和收集卡片，有些页面仍只返回2条或少量岗位。根因可能是游客态只展示少量岗位、需要登录、触发验证码、虚拟列表只保留当前DOM、或分页按钮未暴露。应对策略：记录 collect_round、job_card_count、reader_text_preview，并在报告中说明不绕过登录/验证码限制。"),
        ("3. Qwen抽取阶段超时", "现象：后端已读取到网页文本，但Qwen返回 The read operation timed out。根因是输入文本包含导航、城市筛选和大量无关内容，导致模型处理时间过长。应对策略：增加 focus_job_text 和岗位卡片输入策略，将超时提高到120秒，并优先传岗位正文或卡片。"),
        ("4. Windows浏览器子进程事件循环错误", "现象：Scrapling/Playwright在Windows上启动时报 NotImplementedError。根因是当前 asyncio 事件循环策略不支持 subprocess。应对策略：新增 browser_runtime.ensure_browser_event_loop_policy，在FastAPI、Scrapling和Playwright启动前设置 WindowsProactorEventLoopPolicy。"),
        ("5. 模型与浏览器能力边界不清", "现象：用户期望直接把URL发给模型后模型自行动态渲染，但普通OpenAI-compatible接口不一定提供浏览器工具。根因是模型网页端能力和API能力不等价。应对策略：明确采用“后端浏览器读取 + 大模型理解抽取”的工程化架构，并保留豆包API作为可选对照。"),
    ]:
        add_heading(doc, title, 2)
        add_para(doc, body)

    add_heading(doc, "五、当前系统实现清单", 1)
    add_table(
        doc,
        ["模块", "当前实现", "主要文件"],
        [
            ["前端", "Vue3/Vite页面，支持样例分析、上传分析、采集PoC、AI阅读URL、豆包阅读URL、报告下载", "frontend/src/App.vue; frontend/src/services/api.js"],
            ["后端API", "FastAPI接口，负责pandas分析、报告生成、URL读取、RAG搜索与重建", "backend/app/main.py"],
            ["AI核心", "Qwen API + Tool Calling + RAG上下文 + local_fallback", "backend/app/ai_agent.py"],
            ["复杂URL读取", "腾讯详情适配、Scrapling、Playwright动态渲染、urllib兜底", "backend/app/url_reader.py; backend/app/scraper.py; backend/app/scrapling_reader.py"],
            ["RAG", "ChromaDB向量库、sentence-transformers嵌入、关键词fallback", "backend/app/rag.py; docs/rag_knowledge_base.md"],
            ["启动脚本", "一键启动前后端演示", "start_all.bat; start_demo.bat"],
        ],
        [1.2, 3.6, 1.7],
    )

    add_heading(doc, "六、下一阶段工作重点", 1)
    for item in [
        "建立公开URL回归测试集，分别覆盖腾讯详情页、静态岗位页、BOSS列表页、无薪资页、脚本壳页和错误URL。",
        "将 job_card_count、collect_round、reader_text_preview、reader_errors 等日志在前端做更清晰的分层展示。",
        "继续维护站点级适配器，但避免绕过登录、验证码、付费墙和访问控制。",
        "为AI阅读URL增加批量模式，将多页URL的岗位核心信息合并导出为CSV，再进入pandas统计。",
        "增加模型输出校验器，自动检查salary_text是否来自原文证据，防止大模型臆造薪资。",
        "对RAG引入相关性阈值和rerank，减少弱相关知识片段对报告的干扰。",
        "评估豆包API是否正式支持网页阅读工具；若支持，则将其作为独立对照链路保留。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "七、安装与运行说明", 1)
    add_para(
        doc,
        "当前系统运行需要 Python 后端环境、Node.js 前端环境，以及可选的浏览器自动化依赖。Scrapling已用于复杂URL增强，推荐安装 scrapling[fetchers]；若需要命令行调试，可选安装 scrapling[shell]。不建议在没有明确需求时安装 scrapling[all]，以避免依赖过多和冲突。",
    )
    add_table(
        doc,
        ["依赖/软件", "是否必需", "用途"],
        [
            ["Python 3.11/3.12", "必需", "运行FastAPI、pandas、RAG和抓取模块"],
            ["Node.js LTS", "必需", "运行Vue3/Vite前端"],
            ["playwright chromium", "必需/强烈建议", "动态渲染公开招聘页面"],
            ["scrapling[fetchers]", "建议", "增强复杂URL抓取能力"],
            ["scrapling[shell]", "可选", "命令行调试网页抓取"],
            ["scrapling[ai]/[all]", "暂不必需", "MCP/AI集成或全功能研究场景"],
            ["LibreOffice", "可选", "Word报告渲染为PDF/PNG进行视觉检查"],
        ],
        [1.8, 1.1, 3.6],
    )

    add_para(doc, "日期：2026年5月31日")
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
